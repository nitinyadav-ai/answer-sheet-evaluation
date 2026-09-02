import os
import sys
import json
import warnings
import re
import hashlib
# Suppress all warnings to keep stdout clean for JSON
warnings.filterwarnings("ignore")

import docx
import PyPDF2
from pathlib import Path

from llm_client import generate, strip_reasoning
import parallel_parse as pp


def _load_project_env():
    """Load the project .env into os.environ (without overriding already-set vars) so the provider
    and model settings (LLM_BASE_URL, LLM_API_KEY, KEY_PARSER_MODEL)
    take effect even when this parser is launched as a subprocess that did not inherit them (the
    Flask app invokes it without passing env=)."""
    env_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "..", ".env")
    try:
        with open(env_path) as f:
            for line in f:
                line = line.strip()
                if "=" in line and not line.startswith("#"):
                    k, v = line.split("=", 1)
                    v = v.strip()
                    if v[:1] not in ('"', "'"):
                        v = re.sub(r'\s+#.*$', '', v).strip()  # drop an inline comment from unquoted values
                    os.environ.setdefault(k.strip(), v.strip('"').strip("'"))
    except FileNotFoundError:
        pass


def _sanitize_json_escapes(s):
    """Repair invalid backslash escapes in model JSON so a LaTeX-heavy answer key (\\frac, \\sqrt)
    parses instead of raising 'Invalid \\escape'. Doubles every backslash that is not a genuine JSON
    escape (so LaTeX is preserved as a literal backslash), keeping \\" \\\\ \\/ \\uXXXX; \\n always
    (newlines are common); and \\t \\f \\b \\r only when NOT followed by a letter."""
    out = []
    i, n = 0, len(s)
    while i < n:
        c = s[i]
        if c != '\\':
            out.append(c); i += 1; continue
        nxt = s[i + 1] if i + 1 < n else ''
        if nxt in '"\\/':
            out.append('\\' + nxt); i += 2
        elif nxt == 'u' and re.fullmatch(r'[0-9a-fA-F]{4}', s[i + 2:i + 6] or ''):
            out.append(s[i:i + 6]); i += 6
        elif nxt == 'n':
            out.append('\\n'); i += 2
        elif nxt in 'tfbr' and not (i + 2 < n and s[i + 2].isalpha()):
            out.append('\\' + nxt); i += 2
        else:
            out.append('\\\\'); i += 1
    return ''.join(out)

def extract_text_from_docx(path):
    doc = docx.Document(path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                full_text.append(cell.text)
    return "\n".join(full_text)

def extract_text_from_pdf(path):
    text = ""
    with open(path, 'rb') as f:
        reader = PyPDF2.PdfReader(f)
        for page in reader.pages:
            text += page.extract_text() + "\n"
    return text

def parse_with_gemini(text):
    # Make .env settings (provider + model + keys) visible whether run standalone or from Flask.
    _load_project_env()
    # Parsing the key is the highest-stakes extraction (a weak model once dropped ~72% of a real
    # key's questions, collapsing scores). Default to a capable Qwen3-VL Instruct model; override via
    # KEY_PARSER_MODEL only after A/B-confirming the new model still totals the full marks.
    model_id = os.environ.get("KEY_PARSER_MODEL", "qwen/qwen3-vl-30b-a3b-instruct")

    prompt = f"""
    Extract the exam metadata AND the question/answer information from the following text into a single structured JSON object.

    Return a JSON object with EXACTLY two top-level keys:
    1. "metadata": an object with:
       - class: string (the class / grade this paper is for, e.g. "Class X", "Class XII", "Grade 10". Use an empty string "" if it is not stated in the text.)
       - subject: string (the subject, e.g. "Science", "Computer Science". Use an empty string "" if not stated.)
       - choice_groups: a list describing an INTERNAL CHOICE where the student answers only ONE of
         two-or-more alternatives joined by "OR". This covers BOTH a choice between separately
         numbered questions ("Q31 OR Q32") AND a single question that prints its alternatives as
         sub-parts ("Q31(a) ... OR ... Q31(b)", or "22. (a) ... OR (b) ..."). See rule (A) below for
         the exact splitting + marks handling. For each such choice emit ONE object:
         {{"parent": "<short base id, e.g. Q31>", "members": ["Q31(a)", "Q31(b)"], "required": 1}}.
         "members" MUST be the EXACT question IDs you used as keys in "questions" (so they can be
         matched later). "required" is how many the student must answer (almost always 1). Use [] if
         the paper has no such choices.
       - inline_choice_ids: a list of question IDs whose SINGLE entry already contains an internal
         "OR" between alternatives inside that one question (e.g. one question worth 4 marks whose
         text/answer reads "(iii)(a) ... OR (iii)(b) ..."). These are NOT split into separate entries
         — the whole choice lives in one question. Use [] if there are none.
    2. "questions": a dictionary where keys are question IDs (e.g., Q1, Q2, or specific IDs like SCI10_Q1 if present).
       Each value is a dictionary with the following fields:
       - question_id: string
       - question: string — set this to the EMPTY string "". Do NOT reproduce the question text here:
         the question paper supplies it downstream, so emitting it only wastes output. (The ANSWER field
         below still carries everything the grader needs, INCLUDING any MCQ option text.)
       - answer: string (the correct answer / marking scheme, verbatim).
         For EVERY objective-type question (MCQ, Assertion-Reason, True/False, "choose the correct
         option", match-the-column), you MUST keep BOTH the option identifier AND its full option
         text EXACTLY as printed in the scheme — e.g. "(B) 5", "(D) Neither prime nor composite",
         "(iv) 4/√3". NEVER shorten it to only the value ("5") or only the letter ("B"): the grader
         can match the student on EITHER the identifier OR the text, so both must survive whenever
         the scheme prints them. For subjective questions, give the full marking-scheme answer.
         When the answer contains a PROGRAM, PSEUDOCODE, or an SQL query, wrap that code in a
         [CODE: ... ] fence and reproduce it verbatim — keep every underscore in identifiers
         (product_list, emp_id) as '_', preserve indentation, brackets, colons and operators, and do
         NOT normalise them. This keeps code from being rendered as mathematical notation downstream.
       - type: string (one of: MCQ, Short Answer, Long Answer, Numerical)
       - subject: string
       - marks: number (maximum marks for this question)

    HOW TO HANDLE CHOICES & MULTI-PART QUESTIONS — classify EVERY question that contains an "OR" or
    internal part labels using the rule below. Apply it GENERALLY, regardless of HOW the scheme labels
    the parts ((a)/(b), (i)/(ii)/(iii), roman numerals, bullets, separate question numbers, or a bare
    "OR"), and regardless of whether the parts are printed inline on one line or stacked over lines.
    The DECIDING TEST is the marks semantics — "answer only one" vs "answer all and add up":

    (A) WHOLE-QUESTION CHOICE — the alternatives are joined by "OR" and the student answers only the
        required number (almost always ONE). Each alternative is INDEPENDENTLY worth the SAME full
        marks; the marks are NOT shared and NOT added across them. This holds EVEN when the
        alternatives are printed like sub-parts of one question (e.g. "22. (a) ... OR (b) ...").
          • Emit EACH alternative as its OWN "questions" entry, keyed by the question's base number
            plus a distinct suffix — e.g. "Q22(a)" and "Q22(b)" (mirroring "Q31(a)"/"Q31(b)").
          • Set EACH alternative's "marks" to the marks for answering THAT ONE alternative — i.e. the
            question's full marks (a 2-mark choice => "marks": 2 on EACH entry). DO NOT divide the
            marks between the alternatives and NEVER ADD THEM TOGETHER: two 3-mark alternatives must be
            "marks": 3 and "marks": 3 — NEVER a single merged entry worth 6.
          • ALSO add one object to metadata.choice_groups:
            {{"parent": "Q22", "members": ["Q22(a)", "Q22(b)"], "required": 1}} — "members" being the
            EXACT ids you created above.
          • NEVER collapse the alternatives into one combined entry, and never sum their marks.

    (B) ADDITIVE MULTI-PART — the question has parts the student must ALL answer, whose marks ADD UP
        (e.g. a 5-mark question split (i) 2 + (ii) 3, or a case study (i),(ii),(iii)). The parts are
        NOT joined by "OR".
          • Emit each part as its own entry (e.g. "Q36(i)","Q36(ii)","Q36(iii)") carrying that part's
            own marks. Do NOT list these in choice_groups — they are not a choice.

    (C) ADDITIVE MULTI-PART WITH AN INTERNAL "OR" IN ONE PART — a category-(B) question where a SINGLE
        part offers its own "OR" alternatives (e.g. parts (i),(ii) plus "(iii)(a) ... OR (iii)(b) ..."):
          • Keep it as ONE additive question (do NOT convert the whole question into a choice); its
            marks stay the full additive total.
          • Emit the additive parts as their own entries AND fold the internal OR into the part that
            offers it. Add the question's id to metadata.inline_choice_ids, so the grader accepts
            either alternative for that one part while still summing the other parts.

        WORKED EXAMPLE (category C — a case study; do NOT drop the additive parts). Given a 4-mark
        question printed as:
          "37. (a) <balanced equation> [1]   (b) <two uses> [1]
               (c)(i) <identify A,B + equation> [2]   OR   (c)(ii) <definition + examples> [2]"
        the CORRECT output is THREE entries that SUM to 4 (parts (a)+(b)+(c) = 1+1+2):
          "Q37(a)": {{"marks": 1, "answer": "<equation>", ...}},
          "Q37(b)": {{"marks": 1, "answer": "<two uses>", ...}},
          "Q37(c)": {{"marks": 2, "answer": "(i) <...>  OR  (ii) <...>", ...}}
        and "Q37" added to metadata.inline_choice_ids (the OR lives INSIDE part (c); (a)+(b) are still
        summed). It is WRONG to emit only "Q37(c)(i)"/"Q37(c)(ii)" as a choice_group worth 2 — that
        silently DROPS parts (a),(b) and marks the question 2 instead of 4.

    Hard rules for choices:
    - A question id may appear in AT MOST ONE of choice_groups / inline_choice_ids (never both).
    - The recorded marks must reflect the student answering the REQUIRED number of alternatives (one),
      never all of them: the paper's grand total must NOT inflate just because a choice was offered.
    - NEVER DROP A PART. Every lettered/numbered part that carries its own marks ((a),(b),(i),(ii),...)
      MUST appear in the output with its marks. In particular, when a question has additive parts AND a
      later part offers an internal "OR" (category C above), keep the additive parts — do NOT collapse
      the whole question down to just the OR alternatives. As a self-check: for EACH question, the sum
      of its recorded part-marks (counting an internal OR ONCE, not both alternatives) MUST equal the
      question's printed total marks.

    Text to parse:
    {text}

    Return ONLY the raw JSON. No markdown blocks.
    """
    
    # A full answer key (30-40 questions) overflows the default output cap, silently truncating the
    # JSON. Raise the ceiling + force JSON mode. thinking_budget=0 keeps this verbatim EXTRACTION
    # (not reasoning) so the whole budget goes to output -- a Gemini-only knob, ignored by Qwen.
    # strip_reasoning drops any <think> block a Qwen -Thinking model would emit before the JSON.
    text_out, _in_tok, _out_tok = generate(
        model=model_id, prompt=prompt, temperature=0.0,
        # 32768 is ample for a 40-50 question key (~20k of JSON); env-tunable. OpenRouter
        # PRE-AUTHORISES max_tokens against your balance, so an oversized cap (the old 65536) is
        # rejected with HTTP 402 on a tight/empty balance -- keep this modest.
        max_tokens=int(os.environ.get("KEY_PARSER_MAX_TOKENS", "32768")),
        json_mode=True, thinking_budget=0,
    )
    content = strip_reasoning((text_out or "").strip())

    # JSON mode returns clean JSON, but a LaTeX-heavy key can carry UNescaped backslashes (\frac,
    # \sqrt) -> "Invalid \escape". Try the raw text, then a backslash-repaired copy, each with a
    # brace-extraction fallback for any leading/trailing noise.
    for candidate in (content, _sanitize_json_escapes(content)):
        try:
            return json.loads(candidate)
        except json.JSONDecodeError:
            m = re.search(r'(\{.*\})', candidate, re.DOTALL)
            if m:
                try:
                    return json.loads(m.group(1))
                except json.JSONDecodeError:
                    continue
    raise Exception(f"AI response was not valid JSON (length {len(content)}): {content[:120]}...")


# ---------------------------------------------------------------------------------------------------
# PER-PAGE PARALLEL PATH (default for multi-page PDFs). Each page's answers are extracted concurrently
# by a small, fast call; ONE global pass then detects metadata + internal choices (which need the whole
# paper). Lets the accurate 235B model be used without the whole-document slowness / empty-output.
# ---------------------------------------------------------------------------------------------------
KEY_PER_PAGE_PROMPT = """Extract the questions and their MARKING-SCHEME answers that appear on THIS ONE PAGE of an exam answer key into a JSON object.

Return ONLY: {"questions": {"<id>": {"question_id": "...", "answer": "...", "marks": <number>, "type": "...", "subject": "..."}, ...}}

Rules for THIS PAGE ONLY:
- One entry per question or sub-part printed here, keyed by its id EXACTLY as printed (Q1, Q31(a), Q36(iii)(a), ...).
- "answer": the correct answer / marking scheme for that id, verbatim. For EVERY objective question (MCQ, Assertion-Reason, True/False, match-the-column) keep BOTH the option identifier AND its full option text (e.g. "(B) 5", "(D) Neither prime nor composite") -- never shorten to just the letter or just the value. When the answer is a PROGRAM, PSEUDOCODE or SQL query, wrap that code in a [CODE: ... ] fence and keep it verbatim -- underscores in identifiers stay '_' (product_list, emp_id), indentation/brackets/colons preserved -- so it is not rendered as math downstream.
- "marks": the marks for that id as printed (a number); use 0 if no marks are shown here.
- "type": one of MCQ, Short Answer, Long Answer, Numerical.
- Do NOT analyse "OR"/choices here, and do NOT merge or split parts across pages -- capture exactly what is on THIS page. If the page has no questions (a cover sheet, general instructions, or blank), return {"questions": {}}.

THIS PAGE:
{page_text}

Return ONLY the raw JSON. No markdown."""


# Global metadata + choice pass. Sees the whole paper; returns ONLY metadata (no question list).
KEY_GLOBAL_PROMPT = """You are given the FULL text of an exam answer key. Return ONLY the exam METADATA and its internal-CHOICE structure. Do NOT list the questions.

Return ONLY: {"metadata": {"class": "<e.g. Class X or empty>", "subject": "<e.g. Mathematics or empty>", "choice_groups": [...], "inline_choice_ids": [...]}}

An internal CHOICE is where the student answers only ONE of two-or-more alternatives joined by "OR".
  - choice_groups: for each such choice emit {"parent": "<base id e.g. Q34>", "members": ["Q34(a)", "Q34(b)"], "required": 1}. This covers BOTH a choice between numbered questions ("Q31 OR Q32") AND a choice printed as sub-parts ("34. (a) ... OR (b) ...") AND an OR inside ONE sub-part of a multi-part question (a case study whose part (iii) reads "(iii)(a) ... OR (iii)(b) ..." -> members ["Q36(iii)(a)", "Q36(iii)(b)"]). "members" are the exact printed ids of the alternatives.
  - inline_choice_ids: question ids whose SINGLE entry already contains the whole "OR" inside its own text (not split into separate member ids). Use [] if none.
Rules: a question id appears in AT MOST ONE of choice_groups / inline_choice_ids. Every genuine "OR" in the paper MUST appear as a choice_group (or inline id) so the marks are not double-counted. Use [] when the paper has no choices.

FULL ANSWER KEY TEXT:
{full_text}

Return ONLY the raw JSON. No markdown."""


def parse_key_parallel(page_texts, full_text):
    """Per-page parallel extraction + one global metadata/choice pass -> {metadata, questions}."""
    _load_project_env()
    model_id = os.environ.get("KEY_PARSER_MODEL", "qwen/qwen3-vl-30b-a3b-instruct")
    per_page_max = int(os.environ.get("KEY_PARSER_PAGE_MAX_TOKENS", "8192"))
    global_max = int(os.environ.get("KEY_PARSER_GLOBAL_MAX_TOKENS", "4096"))
    questions, _i1, _o1 = pp.extract_pages_parallel(page_texts, KEY_PER_PAGE_PROMPT, model_id, per_page_max)
    meta, _i2, _o2 = pp.global_metadata_pass(full_text, KEY_GLOBAL_PROMPT, model_id, global_max)
    meta.setdefault("choice_groups", [])
    meta.setdefault("inline_choice_ids", [])
    return {"metadata": meta, "questions": questions}


# --- Parse cache: the key is parsed ONCE per exam but re-uploading the SAME file re-runs the (slow,
# single-call) LLM parse. Cache the parsed JSON by a hash of the raw text + model + prompt version, so an
# identical key returns instantly. Identical input -> identical output, so this can never change a parse;
# BUMP _CACHE_VERSION whenever the parse prompt/schema changes (it invalidates every stale entry). Best-
# effort: any cache read/write error silently falls back to a fresh parse. ------------------------------
_CACHE_VERSION = "2"     # v2: 'question' field emitted empty (question paper supplies the text)


def _cache_root():
    return os.path.join(os.path.dirname(os.path.abspath(__file__)), "..", ".parse_cache")


def _cache_path(raw_text):
    """Absolute path of the cache file for this key text, or None if the cache dir isn't usable."""
    cache_dir = _cache_root()
    try:
        os.makedirs(cache_dir, exist_ok=True)
    except OSError:
        return None
    model = os.environ.get("KEY_PARSER_MODEL", "qwen/qwen3-vl-30b-a3b-instruct")
    digest = hashlib.sha256(f"{_CACHE_VERSION}\x00{model}\x00{raw_text}".encode("utf-8")).hexdigest()
    return os.path.join(cache_dir, f"key_{digest}.json")


def _cache_get(raw_text):
    p = _cache_path(raw_text)
    if p and os.path.exists(p):
        try:
            with open(p) as f:
                return json.load(f)
        except (OSError, ValueError):
            return None
    return None


def _cache_put(raw_text, parsed):
    p = _cache_path(raw_text)
    if not p:
        return
    try:
        with open(p, "w") as f:
            json.dump(parsed, f)
    except OSError:
        pass


def _parse_cached(raw_text):
    """parse_with_gemini, but served from the content-hash cache on a repeat of the SAME key."""
    _load_project_env()          # ensure KEY_PARSER_MODEL (.env) is in the hash BEFORE lookup, so a
                                 # model change correctly invalidates the cache (parse_with_gemini re-loads)
    parsed = _cache_get(raw_text)
    if parsed is not None:
        return parsed
    parsed = parse_with_gemini(raw_text)
    _cache_put(raw_text, parsed)
    return parsed


def main():
    if len(sys.argv) < 2:
        print("Usage: python3 extract_json_from_key.py <file_path>")
        sys.exit(1)
        
    file_path = sys.argv[1]
    ext = Path(file_path).suffix.lower()
    
    try:
        if ext == '.json':
            with open(file_path, 'r') as f:
                content = f.read()
                json.loads(content)          # validate, then echo a raw .json upload through unchanged
                print(content)
                return
        elif ext == '.docx':
            raw_text = extract_text_from_docx(file_path)
            if not raw_text.strip():
                print("ERROR: No text extracted from file.")
                sys.exit(1)
            parsed_json = _parse_cached(raw_text)       # docx has no page structure -> single call (cached)
        elif ext == '.pdf':
            raw_text = "\n".join(pp.pdf_page_texts(file_path))
            if not raw_text.strip():
                print("ERROR: No text extracted from file.")
                sys.exit(1)
            # The answer KEY stays SINGLE-CALL. Per-page parsing was tested and DROPPED marking-scheme
            # marks (which sit in a right-hand column, so a page fragment loses the answer<->marks
            # association: 24/56 subjective answers came back with 0 marks -> total 51 not 80). Accuracy
            # beats the speed gain here, and the key is parsed ONCE per exam at upload. parse_key_parallel
            # (below) is kept for reference / future layout-robust work but is intentionally not routed to.
            parsed_json = _parse_cached(raw_text)       # single-call parse, served from cache on a repeat
        else:
            print(f"ERROR: Unsupported file extension {ext}")
            sys.exit(1)

        print(json.dumps(parsed_json, indent=2))

    except Exception as e:
        print(f"ERROR: {str(e)}", file=sys.stderr)
        sys.exit(1)

if __name__ == "__main__":
    main()

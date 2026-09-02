"""Pre-upload structural validation for question papers and answer keys.

Turns the teacher upload guidelines (docs/UPLOAD_GUIDELINES.md) into automated gates that run the
moment a file is uploaded -- BEFORE the (slow, paid) LLM parse and BEFORE evaluation -- so the input
causes of parse errors are caught up front instead of silently mis-scoring a paper.

Three layers:
  * RAW  (pre-parse):  detects a scanned/image-only or empty file -- the #1 failure, since the key/QP
    parsers read the embedded TEXT layer (PyPDF2/python-docx), never OCR. A scan yields no text.
  * PARSED (post-parse): no questions extracted, or questions with no marks.
  * CROSS-CHECK: compares the answer key against the independently-parsed question paper -- grand-total
    mismatch, and per-question shortfall / inflation / dropped / unknown. Uses the SAME base-number
    logic as the grading-time reconciler (full_evaluator.reconcile_marks_with_question_paper) so the
    two always agree; this just runs it earlier, read-only, so the teacher can fix the upload.

Issues carry a severity: ERROR blocks, WARNING is surfaced but allowed (the grader's reconciler still
corrects the total and flags the affected answers at grading time). Offline except a local file read.
"""
import os
import re

from marks_policy import MARK_STEP, is_valid_mark, quantize_mark

# Same canonicalisation the grading pipeline + reconciler use, so pre-flight groupings are identical.
# full_evaluator imports only qid_utils at module load (no network), so this is safe offline.
try:
    from full_evaluator import (normalize_qid, _base_qnum, effective_choice_marks, subpart_of,
                                _is_under)
except Exception:  # pragma: no cover - fallback keeps validation usable even if the import path shifts
    def normalize_qid(q):
        return str(q)

    def _base_qnum(q):
        m = re.search(r'Q\s*0*(\d+)', str(q), re.IGNORECASE)
        return m.group(1) if m else None

    def effective_choice_marks(leaf_marks, members):
        return None

    def _is_under(entry_id, member_id):
        e, m = str(entry_id), str(member_id)
        return e == m or e.startswith(m + "(") or e.startswith(m + ".")

    def subpart_of(q):
        s = re.sub(r'^\s*(?:Q|Ques|Question|Ans|Answer)?\s*0*\d+', '', str(q), flags=re.IGNORECASE).strip()
        return s if re.search(r'[A-Za-z0-9]', s) else ""

ERROR = "error"
WARNING = "warning"


def _issue(severity, code, message):
    return {"severity": severity, "code": code, "message": message}


def has_blocking(issues):
    """True if any issue is ERROR severity (evaluation / parse should be gated)."""
    return any(i.get("severity") == ERROR for i in (issues or []))


def _fmt(x):
    try:
        f = float(x)
    except (TypeError, ValueError):
        return str(x)
    return str(int(f)) if f == int(f) else str(f)


def _base(qid):
    """Base question number as the reconciler sees it: normalise first ('37'->'Q37'), then extract."""
    return _base_qnum(normalize_qid(qid))


def _to_marks(v):
    """Best-effort marks -> float, or None when there is no usable number."""
    if v is None:
        return None
    if isinstance(v, bool):
        return None
    if isinstance(v, (int, float)):
        return float(v) if v == v else None  # NaN guard
    s = str(v).strip()
    try:
        return float(s)
    except ValueError:
        m = re.search(r'\d+(?:\.\d+)?', s)
        return float(m.group(0)) if m else None


def _questions(obj):
    """The flat {qid: {...}} question map, tolerant of a {metadata, questions} wrapper."""
    if isinstance(obj, dict) and isinstance(obj.get("questions"), dict):
        obj = obj["questions"]
    if not isinstance(obj, dict):
        return {}
    return {k: v for k, v in obj.items() if k not in ("_instructions_", "metadata")}


# --------------------------------------------------------------------------------------------------
# RAW file (pre-parse): scan / no-text-layer detection
# --------------------------------------------------------------------------------------------------
def _extract_raw_text(file_path):
    """(text, page_count) from a PDF/DOCX using the SAME extractors the parsers use (text layer only).
    Raises on a corrupt file; returns ('', 0) for an unhandled extension."""
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".pdf":
        import PyPDF2
        parts = []
        with open(file_path, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            pages = len(reader.pages)
            for pg in reader.pages:
                try:
                    parts.append(pg.extract_text() or "")
                except Exception:
                    pass
        return "\n".join(parts), pages
    if ext == ".docx":
        import docx
        doc = docx.Document(file_path)
        parts = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    parts.append(cell.text)
        return "\n".join(parts), len(parts)
    return "", 0


def validate_raw_file(file_path, kind="answer key"):
    """Pre-parse gate. Detects the file that a text parser cannot read: missing, wrong type, or a
    scanned/image-only PDF/DOCX with no text layer (the biggest single cause of a broken parse)."""
    if not file_path or not os.path.exists(file_path):
        return [_issue(ERROR, "missing_file", f"The {kind} file could not be found on the server.")]
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".json":
        return []  # already-structured upload; nothing to extract
    if ext not in (".pdf", ".docx"):
        return [_issue(ERROR, "unsupported_type",
                       f"Unsupported {kind} file type '{ext or 'unknown'}'. Upload a text-based PDF or a DOCX.")]
    try:
        text, _pages = _extract_raw_text(file_path)
    except Exception as e:
        return [_issue(ERROR, "unreadable",
                       f"Could not read the {kind} file ({e}). Upload a valid text-based PDF or DOCX.")]
    n = len((text or "").strip())
    if n == 0:
        return [_issue(ERROR, "no_text_layer",
                       f"The {kind} has no text layer -- it looks like a scanned image or photo. Upload a "
                       f"text-based PDF or a DOCX. (If you only have a scan, run OCR / 'Make searchable "
                       f"PDF' first, then check the text is selectable.)")]
    if n < 200:
        return [_issue(WARNING, "very_little_text",
                       f"Very little text was extracted from the {kind} ({n} characters) -- it may be a "
                       f"scan or mostly images. Check that the text is selectable in the file.")]
    return []


# --------------------------------------------------------------------------------------------------
# PARSED questions: no questions / missing marks
# --------------------------------------------------------------------------------------------------
def validate_parsed_questions(obj, kind="answer key"):
    issues = []
    qs = _questions(obj)
    if not qs:
        issues.append(_issue(ERROR, "no_questions",
                             f"No questions were extracted from the {kind}. The file may be a scan, empty, "
                             f"or in a layout the parser could not read."))
        return issues
    no_marks = [k for k, v in qs.items()
                if not isinstance(v, dict) or (_to_marks(v.get("marks")) or 0) <= 0]
    if no_marks:
        frac = len(no_marks) / len(qs)
        shown = ", ".join(no_marks[:12]) + (" …" if len(no_marks) > 12 else "")
        sev = ERROR if frac >= 0.5 else WARNING
        issues.append(_issue(sev, "missing_marks",
                             f"{len(no_marks)} of {len(qs)} questions in the {kind} have no marks: {shown}. "
                             f"State the marks next to each question and sub-part."))

    # GRANULARITY: marks are awarded in half-mark steps, so a question whose MAXIMUM is not a multiple
    # of 0.5 can never be scored legally -- every mark on it would be capped at that illegal ceiling.
    # In practice this means the parser misread the marks ("[2]" read as "[0.8]", a merged sub-part
    # total). Reported, never auto-corrected: the answer key is the teacher's ground truth, and quietly
    # rewriting it would hide a real parse error. WARNING, not ERROR -- the run is still gradeable, and
    # the teacher can fix it in the marks editor, which snaps its own inputs to 0.5.
    odd = []
    for k, v in qs.items():
        if not isinstance(v, dict):
            continue
        m = _to_marks(v.get("marks"))
        if m is None or m <= 0 or is_valid_mark(m):
            continue
        odd.append(f"{k} ({_fmt(m)})")
    if odd:
        shown = ", ".join(odd[:12]) + (" …" if len(odd) > 12 else "")
        issues.append(_issue(WARNING, "marks_not_half_step",
                             f"{len(odd)} of {len(qs)} questions in the {kind} carry marks that are not a "
                             f"multiple of {MARK_STEP}: {shown}. Marks are awarded in half-mark steps, so "
                             f"these were most likely misread. Check them in the marks editor."))
    return issues


# --------------------------------------------------------------------------------------------------
# Marks-by-base (shared with the reconciler's model of how merges collapse marks)
# --------------------------------------------------------------------------------------------------
def qp_marks_by_base(qp):
    """{base_number: marks} from the question paper -- MAX across siblings (never over-counts)."""
    out = {}
    for k, v in _questions(qp).items():
        if not isinstance(v, dict):
            continue
        b = _base(k)
        if b is None:
            continue
        m = _to_marks(v.get("marks"))
        if m is None:
            continue
        out[b] = max(out.get(b, 0.0), m)
    return out


def qp_question_by_base(qp):
    """{base_number: question_text} from the question paper, for the marks editor's guided cards.
    Prefers the entry whose qid IS the base (the stem, no sub-part token) over a sub-part, and the
    longer text within a tier. The paper's `question` is clean, unlike the answer key's own
    `question` field which is polluted to the answer for objective questions (key Q1 question='True')."""
    best = {}
    for k, v in _questions(qp).items():
        if not isinstance(v, dict):
            continue
        b = _base(k)
        if b is None:
            continue
        q = str(v.get("question") or "").strip()
        if not q:
            continue
        is_stem = not subpart_of(normalize_qid(k))
        prev = best.get(b)
        if (prev is None
                or (is_stem and not prev["stem"])
                or (is_stem == prev["stem"] and len(q) > len(prev["q"]))):
            best[b] = {"q": q, "stem": is_stem}
    return {b: d["q"] for b, d in best.items()}


def _leaves_by_base(key):
    """{base_number: {leaf_qid: marks}} over every answer-key entry with usable marks. Shared by the
    effective-marks calc and the choice-group suggester so both reason over the SAME leaves."""
    leaves_by_base = {}
    for k, v in _questions(key).items():
        if not isinstance(v, dict):
            continue
        b = _base(k)
        if b is None:
            continue
        m = _to_marks(v.get("marks"))
        if m is None:
            continue
        leaves_by_base.setdefault(b, {})[k] = m
    return leaves_by_base


def key_effective_marks_by_base(key, choices):
    """{base_number: effective marks} for the answer key AFTER applying choices, using the SAME model
    as the pipeline's choice merge (full_evaluator.effective_choice_marks): for a base offering an
    'answer any one' choice, effective = sum(shared additive parts) + MAX over alternatives (sum of
    that alternative's sub-parts); every other base is a plain additive sum. Prefix-tolerant, so a
    member 'Q34(a)' claims 'Q34(a)(i)'... -- matching what grading actually scores."""
    groups = (choices or {}).get("choice_groups") or []
    members_by_base = {}
    for g in groups:
        members = g.get("members") or []
        if not members:
            continue
        b = _base(g.get("parent")) if g.get("parent") else _base(members[0])
        if b is not None:
            members_by_base.setdefault(b, []).extend(members)

    out = {}
    for b, leaves in _leaves_by_base(key).items():
        eff = effective_choice_marks(leaves, members_by_base[b]) if b in members_by_base else None
        out[b] = eff if eff is not None else sum(leaves.values())
    return out


def _top_alternative_id(leaf_qid):
    """Choice-alternative id a leaf belongs to = base + its FIRST sub-part token, e.g. 'Q34(a)(i)' ->
    'Q34(a)', 'Q22(b)' -> 'Q22(b)', 'Q5.a' -> 'Q5.a'. A leaf with no sub-part ('Q5') returns '' (it is
    a shared/common part, not an alternative). This is the granularity `effective_choice_marks` groups
    at via `_is_under`, so member marks derived here match what grading would score."""
    b = _base(leaf_qid)
    if b is None:
        return ""
    sub = subpart_of(normalize_qid(leaf_qid))
    m = re.match(r'\([^)]*\)|\.[0-9A-Za-z]+|[0-9A-Za-z]+', sub)
    return f"Q{b}{m.group(0)}" if m else ""


def ungrouped_choice_bases(key, choices):
    """Bases that may hide an ungrouped 'answer any one' choice, used to catch a LOST choice sidecar --
    the failure mode where the key parse dropped its `metadata`, so `choice_groups` came back empty and
    every alternative is then counted ADDITIVELY, silently inflating the key's total.

    Fires ONLY when the key declares no choice group at all. A key with any group parsed successfully is
    trusted as-is and returns [], so a working parse is never second-guessed and a base whose sub-parts
    are genuinely additive is not flagged merely for having sub-parts.

    Sub-part COUNT is the whole test -- alternatives being worth the SAME was measured and rejected: a
    case study whose OR sits only in its last part scores [1, 1, 4], so equal-marks missed 3 of 9 real
    choices. Deliberately does not guess a corrected total: with the choice data gone there is no way to
    know which alternative the paper offers, so the caller keeps the additive sum and merely says so."""
    if (choices or {}).get("choice_groups"):
        return []                                    # choices parsed fine -> nothing was lost
    bases = [b for b, leaves in _leaves_by_base(key).items()
             if len({a for a in (_top_alternative_id(lid) for lid in leaves) if a}) >= 2]
    return _sort_bases(bases)


def choices_lost_issues(key, choices):
    """WARNING when the key's total is an unverified additive sum because its choice data is missing.
    Independent of the question paper on purpose: `cross_check` already catches an inflated total when a
    paper is present, but with no paper uploaded nothing else looks at this and the UI would otherwise
    report the inflated sum as verified."""
    bases = ungrouped_choice_bases(key, choices)
    if not bases:
        return []
    shown = ", ".join(f"Q{b}" for b in bases[:8]) + ("..." if len(bases) > 8 else "")
    n = len(bases)
    return [_issue(WARNING, "choices_unavailable",
                   f"No 'answer any one' choices were recorded for this answer key, but {n} question"
                   f"{'' if n == 1 else 's'} ({shown}) have two or more labelled alternatives. If any of "
                   f"those is a choice, its alternatives are being added together, so the key's total is "
                   f"overstated. Re-parse the answer key, or set the choices in 'Review & fix marks'.")]


def _suggest_choice_groups(key, choices, qp):
    """Candidate ungrouped 'answer any one' choices, derived purely from the marks arithmetic (never a
    naive leaf-max). A base is suggested only when grouping its distinct top-level sub-part alternatives
    would drop its effective marks EXACTLY onto the question paper's marks -- the fingerprint of an
    OR-pair the parser laid out as two separately-counted questions. A genuine misread/duplicate won't
    land on the paper's number, so it is never suggested. Bases already carrying a choice-group member
    are left to the full editor. Returns [{base, paper, current_sum, members:[{qid, marks}]}], sorted."""
    qp_by = qp_marks_by_base(qp)
    if not qp_by:
        return []
    key_eff = key_effective_marks_by_base(key, choices)
    grouped_bases = set()
    for g in (choices or {}).get("choice_groups") or []:
        for mem in (g.get("members") or []):
            gb = _base(mem)
            if gb is not None:
                grouped_bases.add(gb)

    out = []
    for b, leaves in _leaves_by_base(key).items():
        qpm = qp_by.get(b)
        if qpm is None or b in grouped_bases:
            continue
        if (key_eff.get(b, 0.0) - qpm) <= 1e-6:          # not inflated -> no double-counted choice here
            continue
        alt_of = {lid: _top_alternative_id(lid) for lid in leaves}
        members = sorted({a for a in alt_of.values() if a})
        if len(members) < 2:
            continue
        eff = effective_choice_marks(leaves, members)
        cur = sum(leaves.values())
        if eff is None or (cur - eff) <= 1e-6 or abs(eff - qpm) > 1e-6:
            continue                                     # grouping must collapse AND land exactly on paper
        out.append({
            "base": b, "paper": qpm, "current_sum": cur,
            "members": [{"qid": a, "marks": sum(mk for lid, mk in leaves.items() if alt_of[lid] == a)}
                        for a in members],
        })
    out.sort(key=lambda r: (0, int(r["base"])) if str(r["base"]).isdigit() else (1, str(r["base"])))
    return out


# --------------------------------------------------------------------------------------------------
# Cross-check: answer key vs question paper
# --------------------------------------------------------------------------------------------------
def _sort_bases(bases):
    return sorted(bases, key=lambda x: (0, int(x)) if str(x).isdigit() else (1, str(x)))


def cross_check(key, choices, qp):
    """Compare the key's effective marks to the paper's, per base question + grand total. All findings
    are WARNINGS -- the grading-time reconciler corrects the total and flags the answers regardless --
    but each almost always means the upload can be fixed."""
    issues = []
    qp_by = qp_marks_by_base(qp)
    if not qp_by:
        return issues  # nothing to check against (a paper with no parseable marks is caught elsewhere)
    key_by = key_effective_marks_by_base(key, choices)

    key_total, qp_total = sum(key_by.values()), sum(qp_by.values())
    if abs(key_total - qp_total) > 1e-6:
        if qp_total > key_total:
            gap = f"short by {_fmt(qp_total - key_total)}"
        else:
            gap = f"over by {_fmt(key_total - qp_total)}"
        issues.append(_issue(WARNING, "total_mismatch",
                             f"The answer key totals {_fmt(key_total)} marks but the question paper totals "
                             f"{_fmt(qp_total)} ({gap}). This usually means the key parse dropped or "
                             f"duplicated part of a question -- please re-check the key. (Grading will "
                             f"correct the total against the paper and flag the affected answers.)"))

    shortfalls, inflations, missing, extra = [], [], [], []
    for b in _sort_bases(set(qp_by) | set(key_by)):
        qpm, km = qp_by.get(b), key_by.get(b)
        if qpm is None:
            extra.append(b)
        elif km is None:
            missing.append((b, qpm))
        elif qpm - km > 1e-6:
            shortfalls.append((b, km, qpm))
        elif km - qpm > 1e-6:
            inflations.append((b, km, qpm))

    def _cap(items, render):
        return ", ".join(render(x) for x in items[:12]) + (" …" if len(items) > 12 else "")

    if missing:
        issues.append(_issue(WARNING, "missing_questions",
                             "In the question paper but MISSING from the answer key: "
                             + _cap(missing, lambda t: f"Q{t[0]} (worth {_fmt(t[1])})")
                             + ". Add these questions (with answers + marks) to the key."))
    if shortfalls:
        issues.append(_issue(WARNING, "under_marked",
                             "Worth fewer marks in the key than in the paper (a part may be dropped): "
                             + _cap(shortfalls, lambda t: f"Q{t[0]} {_fmt(t[1])}/{_fmt(t[2])}")
                             + "."))
    if inflations:
        issues.append(_issue(WARNING, "over_marked",
                             "Worth more marks in the key than in the paper (a choice/duplicate may be "
                             "double-counted): " + _cap(inflations, lambda t: f"Q{t[0]} {_fmt(t[1])}/{_fmt(t[2])}")
                             + "."))
    if extra:
        issues.append(_issue(WARNING, "unknown_questions",
                             "In the answer key but not the question paper: "
                             + _cap(extra, lambda b: f"Q{b}")
                             + ". The paper's parse may have dropped these (e.g. an objective section), "
                               "or the numbering differs -- re-check that the paper was fully captured."))
    return issues


def compute_marks_mismatch(key, choices, qp):
    """Structured comparison of the key's effective marks vs the paper's, for the teacher 'marks
    source' chooser. Returns:
        {mismatch: bool, key_total, qp_total, recommended: 'question_paper'|'answer_key',
         per_question: [{q: 'Q34', key: <or None>, qp: <or None>}, ...]}  (only differing questions)
    'recommended' defaults to the QUESTION PAPER (structurally simpler, the native marks skeleton),
    unless the paper itself has no usable marks -- then the key is the better bet."""
    key_by = key_effective_marks_by_base(key, choices)
    qp_by = qp_marks_by_base(qp)
    per_q = []
    for b in _sort_bases(set(key_by) | set(qp_by)):
        km, qm = key_by.get(b), qp_by.get(b)
        if km is None or qm is None or abs(km - qm) > 1e-6:
            per_q.append({"q": f"Q{b}", "key": km, "qp": qm})
    key_total, qp_total = sum(key_by.values()), sum(qp_by.values())
    mismatch = bool(per_q) or abs(key_total - qp_total) > 1e-6
    recommended = "answer_key" if not qp_by else "question_paper"
    return {"mismatch": mismatch, "key_total": key_total, "qp_total": qp_total,
            "recommended": recommended, "per_question": per_q[:60]}


def build_marks_breakdown(key, choices, qp):
    """Per-answer-key-entry rows for the teacher's editable marks breakdown: each entry's key marks +
    the question paper's marks for that base question (reference), its choice-group index, and a short
    answer preview. Plus effective totals + the mismatch flag. Rows are sorted by base question."""
    qs = _questions(key)
    qp_by = qp_marks_by_base(qp)
    qp_q = qp_question_by_base(qp)
    groups = (choices or {}).get("choice_groups") or []
    group_members = [(str(m), gi) for gi, g in enumerate(groups) for m in (g.get("members") or [])]

    def _group_for(qid):
        """Which choice group a row belongs to. Exact match first (unchanged behaviour), then the
        member the row sits UNDER: a saved group's members are the ALTERNATIVE ids (Q32(A)), so its
        finer leaves (Q32(A)(I)...) must still render as grouped -- otherwise the teacher's saved
        choice looks dissolved the moment the panel reloads."""
        s = str(qid)
        for m, gi in group_members:
            if s == m:
                return gi
        for m, gi in group_members:
            if _is_under(s, m):
                return gi
        return None

    rows = []
    for qid, v in qs.items():
        if not isinstance(v, dict):
            continue
        b = _base(qid)
        rows.append({
            "qid": str(qid),
            "base": ("" if b is None else str(b)),
            "key_marks": _to_marks(v.get("marks")),
            "qp_marks": (qp_by.get(b) if b is not None else None),
            "qp_question": ((qp_q.get(b) or "")[:300] if b is not None else ""),
            "group": _group_for(qid),
            "answer_preview": (str(v.get("answer") or "")[:80]),
        })
    rows.sort(key=lambda r: ((0, int(r["base"]), r["qid"]) if r["base"].isdigit() else (1, 0, r["qid"])))
    key_eff = key_effective_marks_by_base(key, choices)
    mm = compute_marks_mismatch(key, choices, qp) if qp_by else {"mismatch": False}
    return {
        "rows": rows,
        "groups": [[str(m) for m in (g.get("members") or [])] for g in groups],
        "key_total": sum(key_eff.values()),
        "qp_total": sum(qp_by.values()),
        "mismatch": mm.get("mismatch", False),
        "suggested_groups": _suggest_choice_groups(key, choices, qp),
        # Bases that may hide a choice the parse lost -> key_total above is an unverified additive sum,
        # so the UI must not announce it as verified. Empty on every key whose choices parsed.
        "choices_missing": ungrouped_choice_bases(key, choices),
    }


def apply_marks_corrections(key, choices, corr):
    """Apply the teacher's edited breakdown to the answer key + choices sidecar. Returns
    (new_questions_dict, new_choices_dict). corr = {marks:{qid:n}, added:[{q,marks,answer?}],
    removed:[qid], choice_groups:[[qid,...]]}. An added entry WITHOUT an answer is flagged
    key_parse_missing (manual grading), mirroring the reconciler's injected-question shape so grading
    surfaces it for review instead of auto-scoring it against an empty expected answer."""
    corr = corr or {}
    qs = {k: (dict(v) if isinstance(v, dict) else v) for k, v in _questions(key).items()}

    for qid in (corr.get("removed") or []):
        qs.pop(str(qid), None)

    for qid, m in (corr.get("marks") or {}).items():
        mv = _to_marks(m)
        if mv is not None and isinstance(qs.get(str(qid)), dict):
            # Snap to the half-mark ladder the editor's inputs already advertise (step="0.5"). This is
            # the teacher TYPING a value, not a parsed key being rewritten behind their back -- so
            # correcting it here is what the UI promises. A parsed non-half maximum is only FLAGGED
            # (see _marks_granularity_issues), never silently changed.
            qs[str(qid)]["marks"] = quantize_mark(mv)

    subj = next((v.get("subject") for v in qs.values()
                 if isinstance(v, dict) and v.get("subject")), "")
    for a in (corr.get("added") or []):
        raw = str((a or {}).get("q") or "").strip()
        if not raw:
            continue
        qid = normalize_qid(raw)
        mv = quantize_mark(_to_marks((a or {}).get("marks")) or 0.0)
        ans = str((a or {}).get("answer") or "").strip()
        entry = {"question_id": qid, "answer": ans, "marks": mv, "type": "",
                 "subject": subj, "teacher_added": True}
        if not ans:
            entry["key_parse_missing"] = True
            entry["reconcile_note"] = ("Added by the teacher; no expected answer was provided, so grade "
                                       "the student's response to this question manually.")
        qs[qid] = entry

    new_choices = dict(choices or {})
    new_choices.setdefault("inline_choice_ids", (choices or {}).get("inline_choice_ids") or [])
    groups = corr.get("choice_groups")
    if groups is not None:
        new_choices["choice_groups"] = [_editor_choice_group(g) for g in groups if g]
    else:
        new_choices["choice_groups"] = (choices or {}).get("choice_groups") or []
    return qs, new_choices


def _prefix_at(qid, depth):
    """Base + the FIRST `depth` sub-part tokens: 'Q32(A)(I)' @1 -> 'Q32(A)', @2 -> 'Q32(A)(I)'.
    '' when the id has no base (caller then leaves the selection alone)."""
    b = _base(qid)
    if b is None:
        return ""
    toks = re.findall(r'\([^)]*\)|\.[0-9A-Za-z]+', subpart_of(normalize_qid(qid)))
    return f"Q{b}" + "".join(toks[:depth])


def _collapse_to_alternatives(members):
    """Collapse an editor selection to the real CHOICE ALTERNATIVES.

    The editor tags the ROWS the teacher ticked, so when each alternative is itself multi-part
    (Q32(A)(I..IV) OR Q32(B)(I..IV)) the raw list declares EIGHT mutually-exclusive alternatives and
    effective_choice_marks returns max(1,1,...) = 1 instead of max(4,4) = 4 -- capping the question at
    1 mark. Collapsing here (the single choke point every save passes through) fixes every editor path
    at once and keeps what is PERSISTED and GRADED correct whatever the UI posts.

    Rule: the SHALLOWEST sub-part depth at which the selection splits into >= 2 distinct alternatives.
    Going deeper only when needed is load-bearing -- a choice nested inside a sub-part
    (Q34(IV)(A) OR Q34(IV)(B)) does NOT split at depth 1, and a blanket first-token collapse would fold
    both into 'Q34(IV)' and destroy the group. A selection that never splits is returned unchanged, so
    anything this cannot interpret is passed through rather than silently reshaped."""
    mem = [str(x) for x in members if str(x).strip()]
    if len(mem) < 2:
        return mem
    for depth in (1, 2, 3, 4):
        alts = list(dict.fromkeys(_prefix_at(q, depth) for q in mem))   # first-seen order
        if any(a == "" for a in alts):
            return mem                  # un-parseable id in the selection -> keep exactly as posted
        if len(alts) >= 2:
            return alts
    return mem


def _editor_choice_group(members):
    """Build the sidecar choice-group shape the GRADER expects — {parent, members, required}.
    full_evaluator.merge_choice_groups keys the collapse off `parent`; writing it here (derived from
    the members' shared base, e.g. 'Q22(A)' -> 'Q22') is what makes an editor-grouped choice actually
    collapse to a single counted-once entry at grading time instead of the pair being SUMMED."""
    mem = _collapse_to_alternatives(members)
    base = _base(mem[0]) if mem else None
    return {"parent": (f"Q{base}" if base else None), "members": mem, "required": 1}


# --------------------------------------------------------------------------------------------------
# Top-level orchestrators used by the Flask routes
# --------------------------------------------------------------------------------------------------
def validate_question_paper_structure(qp):
    """Heuristic completeness check on a PARSED question paper. A LEADING gap (numbering that doesn't
    start at Q1) usually means the parser DROPPED an early section -- most often the 1-mark objective
    'Section A' that sits under the cover/instructions on page 1 -- and an INTERNAL gap means a
    question in the middle was missed. Both are WARNINGS (a paper may genuinely start above 1), but
    they almost always mean the parse is incomplete and the paper should be re-uploaded / re-parsed."""
    nums = []
    for q in _questions(qp):
        b = _base(q)
        if b is not None and str(b).isdigit():
            nums.append(int(b))
    issues = []
    if not nums:
        return issues
    bases = set(nums)
    lo, hi = min(bases), max(bases)
    if lo > 1:
        issues.append(_issue(WARNING, "paper_leading_gap",
                             f"The question paper's numbering starts at Q{lo}, not Q1 -- an earlier "
                             f"section (often the 1-mark objective / MCQ 'Section A') may not have been "
                             f"captured. Check that every question was extracted; re-upload the paper if "
                             f"the start is missing."))
    missing = [n for n in range(lo, hi + 1) if n not in bases]
    if missing:
        shown = ", ".join(f"Q{n}" for n in missing[:12]) + (" …" if len(missing) > 12 else "")
        issues.append(_issue(WARNING, "paper_internal_gap",
                             f"Gaps in the parsed paper's numbering ({shown}) -- these questions may have "
                             f"been missed. Check that the paper parsed completely."))
    return issues


def validate_question_paper(raw_path, qp_json):
    """Full check for a freshly-uploaded question paper (raw + parsed)."""
    issues = validate_raw_file(raw_path, "question paper")
    if not has_blocking(issues):
        issues += validate_parsed_questions(qp_json, "question paper")
        issues += validate_question_paper_structure(qp_json)
    return issues


def validate_answer_key(raw_path, key_json, choices_json, qp_json=None):
    """Full check for a freshly-uploaded answer key (raw + parsed + cross-check vs the paper)."""
    issues = validate_raw_file(raw_path, "answer key")
    if not has_blocking(issues):
        issues += validate_parsed_questions(key_json, "answer key")
        issues += choices_lost_issues(key_json, choices_json)
        if qp_json is not None and _questions(qp_json):
            issues += cross_check(key_json, choices_json, qp_json)
        else:
            issues.append(_issue(WARNING, "no_question_paper",
                                 "No question paper has been uploaded yet, so the answer key could not be "
                                 "cross-checked. Upload the question paper (Step 1) to enable the automatic "
                                 "marks check."))
    return issues


def validate_for_evaluation(key_json, choices_json, qp_json):
    """Final gate just before evaluation (post-parse only). Blocks on ERROR; returns all issues."""
    issues = []
    if qp_json is None or not _questions(qp_json):
        issues.append(_issue(ERROR, "no_question_paper",
                             "No question paper is available. Upload and parse the question paper (Step 1) "
                             "before evaluating."))
    issues += validate_parsed_questions(key_json, "answer key")
    issues += choices_lost_issues(key_json, choices_json)
    if qp_json is not None:
        issues += validate_parsed_questions(qp_json, "question paper")
        issues += cross_check(key_json, choices_json, qp_json)
    return issues
